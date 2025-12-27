from io import BytesIO
from typing import Dict, Any
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont
import io

class ExportService:

    def generate_pdf(self, data: Dict[str, Any], doc_type: str) -> BytesIO:
        """Generate PDF from document data"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=TA_CENTER
        )

        title = "INVOICE" if doc_type == 'invoice' else "QUOTATION"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.3*inch))

        # Document info
        doc_number = data.get('invoice_number' if doc_type == 'invoice' else 'quote_number', 'N/A')
        info_data = [
            [f"{title} Number:", doc_number],
            ["Date:", data.get('date', 'N/A')]
        ]

        info_table = Table(info_data, colWidths=[2*inch, 3*inch])
        info_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.grey),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 0.3*inch))

        # Bill To / To section
        bill_label = "Bill To:" if doc_type == 'invoice' else "To:"
        story.append(Paragraph(bill_label, styles['Heading2']))
        customer_info = f"""
        <b>{data.get('customer_name', 'N/A')}</b><br/>
        {data.get('address', 'N/A')}<br/>
        {data.get('city', 'N/A')}, {data.get('country', 'N/A')}
        """
        story.append(Paragraph(customer_info, styles['Normal']))
        story.append(Spacer(1, 0.3*inch))

        # Items table
        items_data = [['Description', 'Quantity', 'Unit Price', 'Amount']]
        for item in data.get('items', []):
            items_data.append([
                item.get('description', ''),
                str(item.get('quantity', 0)),
                f"{data.get('currency', 'NGN')} {item.get('unit_price', 0):,.2f}",
                f"{data.get('currency', 'NGN')} {item.get('amount', 0):,.2f}"
            ])

        items_table = Table(items_data, colWidths=[3*inch, 1*inch, 1.5*inch, 1.5*inch])
        items_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(items_table)
        story.append(Spacer(1, 0.3*inch))

        # Totals
        totals_data = [
            ['Subtotal:', f"{data.get('currency', 'NGN')} {data.get('subtotal', 0):,.2f}"],
            [f"Tax ({data.get('tax_rate', 0)*100:.0f}%):", f"{data.get('currency', 'NGN')} {data.get('tax_amount', 0):,.2f}"]
        ]

        if doc_type == 'invoice':
            totals_data.append([
                f"Delivery ({data.get('delivery_rate', 0)*100:.0f}%):",
                f"{data.get('currency', 'NGN')} {data.get('delivery_amount', 0):,.2f}"
            ])

        totals_data.append(['', ''])  # Spacer
        totals_data.append([
            'TOTAL:',
            f"{data.get('currency', 'NGN')} {data.get('total', 0):,.2f}"
        ])

        totals_table = Table(totals_data, colWidths=[5*inch, 2*inch])
        totals_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, -1), (-1, -1), 14),
            ('LINEABOVE', (0, -1), (-1, -1), 2, colors.black),
            ('TOPPADDING', (0, -1), (-1, -1), 10),
        ]))
        story.append(totals_table)

        doc.build(story)
        buffer.seek(0)
        return buffer

    def generate_docx(self, data: Dict[str, Any], doc_type: str) -> BytesIO:
        """Generate DOCX from document data"""
        doc = Document()

        # Title
        title = "INVOICE" if doc_type == 'invoice' else "QUOTATION"
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Document info
        doc_number = data.get('invoice_number' if doc_type == 'invoice' else 'quote_number', 'N/A')
        p = doc.add_paragraph()
        p.add_run(f"{title} Number: ").bold = True
        p.add_run(doc_number)
        p = doc.add_paragraph()
        p.add_run("Date: ").bold = True
        p.add_run(data.get('date', 'N/A'))

        doc.add_paragraph()

        # Bill To / To
        bill_label = "Bill To:" if doc_type == 'invoice' else "To:"
        doc.add_heading(bill_label, 2)
        doc.add_paragraph(data.get('customer_name', 'N/A')).bold = True
        doc.add_paragraph(data.get('address', 'N/A'))
        doc.add_paragraph(f"{data.get('city', 'N/A')}, {data.get('country', 'N/A')}")

        doc.add_paragraph()

        # Items table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Description'
        hdr_cells[1].text = 'Quantity'
        hdr_cells[2].text = 'Unit Price'
        hdr_cells[3].text = 'Amount'

        for item in data.get('items', []):
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('description', '')
            row_cells[1].text = str(item.get('quantity', 0))
            row_cells[2].text = f"{data.get('currency', 'NGN')} {item.get('unit_price', 0):,.2f}"
            row_cells[3].text = f"{data.get('currency', 'NGN')} {item.get('amount', 0):,.2f}"

        doc.add_paragraph()

        # Totals
        p = doc.add_paragraph()
        p.add_run(f"Subtotal: {data.get('currency', 'NGN')} {data.get('subtotal', 0):,.2f}").alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = doc.add_paragraph()
        p.add_run(f"Tax ({data.get('tax_rate', 0)*100:.0f}%): {data.get('currency', 'NGN')} {data.get('tax_amount', 0):,.2f}")

        if doc_type == 'invoice':
            p = doc.add_paragraph()
            p.add_run(f"Delivery ({data.get('delivery_rate', 0)*100:.0f}%): {data.get('currency', 'NGN')} {data.get('delivery_amount', 0):,.2f}")

        p = doc.add_paragraph()
        run = p.add_run(f"TOTAL: {data.get('currency', 'NGN')} {data.get('total', 0):,.2f}")
        run.bold = True
        run.font.size = Pt(16)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_image(self, data: Dict[str, Any], doc_type: str) -> BytesIO:
        """Generate PNG image from document data"""
        # Create image
        img = Image.new('RGB', (800, 1000), color='white')
        draw = ImageDraw.Draw(img)

        try:
            title_font = ImageFont.truetype("arial.ttf", 32)
            heading_font = ImageFont.truetype("arial.ttf", 16)
            text_font = ImageFont.truetype("arial.ttf", 12)
        except:
            title_font = ImageFont.load_default()
            heading_font = ImageFont.load_default()
            text_font = ImageFont.load_default()

        y = 50

        # Title
        title = "INVOICE" if doc_type == 'invoice' else "QUOTATION"
        draw.text((400, y), title, fill='black', font=title_font, anchor="mm")
        y += 60

        # Document info
        doc_number = data.get('invoice_number' if doc_type == 'invoice' else 'quote_number', 'N/A')
        draw.text((50, y), f"{title} Number: {doc_number}", fill='black', font=text_font)
        y += 25
        draw.text((50, y), f"Date: {data.get('date', 'N/A')}", fill='black', font=text_font)
        y += 40

        # Customer info
        bill_label = "Bill To:" if doc_type == 'invoice' else "To:"
        draw.text((50, y), bill_label, fill='black', font=heading_font)
        y += 25
        draw.text((50, y), data.get('customer_name', 'N/A'), fill='black', font=text_font)
        y += 20
        draw.text((50, y), data.get('address', 'N/A'), fill='black', font=text_font)
        y += 20
        draw.text((50, y), f"{data.get('city', 'N/A')}, {data.get('country', 'N/A')}", fill='black', font=text_font)
        y += 40

        # Items
        draw.text((50, y), "Items:", fill='black', font=heading_font)
        y += 30

        for item in data.get('items', []):
            desc = item.get('description', '')
            qty = item.get('quantity', 0)
            price = item.get('unit_price', 0)
            amount = item.get('amount', 0)
            currency = data.get('currency', 'NGN')

            draw.text((50, y), f"{desc}", fill='black', font=text_font)
            y += 20
            draw.text((70, y), f"Qty: {qty} x {currency} {price:,.2f} = {currency} {amount:,.2f}", fill='gray', font=text_font)
            y += 30

        # Totals
        y += 20
        draw.text((50, y), f"Subtotal: {data.get('currency', 'NGN')} {data.get('subtotal', 0):,.2f}", fill='black', font=text_font)
        y += 25
        draw.text((50, y), f"Tax ({data.get('tax_rate', 0)*100:.0f}%): {data.get('currency', 'NGN')} {data.get('tax_amount', 0):,.2f}", fill='black', font=text_font)
        y += 25

        if doc_type == 'invoice':
            draw.text((50, y), f"Delivery ({data.get('delivery_rate', 0)*100:.0f}%): {data.get('currency', 'NGN')} {data.get('delivery_amount', 0):,.2f}", fill='black', font=text_font)
            y += 25

        y += 10
        draw.text((50, y), f"TOTAL: {data.get('currency', 'NGN')} {data.get('total', 0):,.2f}", fill='black', font=heading_font)

        buffer = BytesIO()
        img.save(buffer, format='PNG')
        buffer.seek(0)
        return buffer
